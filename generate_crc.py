#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tự động tạo thiết kế CRC (Class-Responsibility-Collaboration) từ mã nguồn Java
Yêu cầu: pip install python-docx
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Set
from dataclasses import dataclass, field
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Cảnh báo: python-docx chưa được cài đặt. Chạy: pip install python-docx")

@dataclass
class ClassInfo:
    """Thông tin về một lớp"""
    name: str
    package: str
    type: str  # 'class', 'interface', 'enum'
    attributes: List[str] = field(default_factory=list)
    methods: List[str] = field(default_factory=list)
    imports: List[str] = field(default_factory=list)
    responsibilities: List[str] = field(default_factory=list)
    collaborators: Set[str] = field(default_factory=set)
    file_path: str = ""

@dataclass
class CRCCard:
    """Thẻ CRC cho một lớp"""
    class_name: str
    package: str
    responsibilities: List[str]
    collaborators: List[str]
    attributes: List[str]
    methods: List[str]

class JavaParser:
    """Parser để đọc và phân tích file Java"""
    
    def __init__(self, base_path: str):
        self.base_path = Path(base_path)
        self.classes: Dict[str, ClassInfo] = {}
        
    def parse_file(self, file_path: Path) -> ClassInfo:
        """Parse một file Java và trả về ClassInfo"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Tìm package
        package_match = re.search(r'package\s+([\w.]+);', content)
        package = package_match.group(1) if package_match else ""
        
        # Tìm class/interface/enum
        class_match = re.search(r'public\s+(class|interface|enum)\s+(\w+)', content)
        if not class_match:
            return None
        
        class_type = class_match.group(1)
        class_name = class_match.group(2)
        
        class_info = ClassInfo(
            name=class_name,
            package=package,
            type=class_type,
            file_path=str(file_path.relative_to(self.base_path))
        )
        
        # Tìm imports
        imports = re.findall(r'import\s+([\w.]+);', content)
        class_info.imports = [imp for imp in imports if 'com.bluemoon' in imp]
        
        # Tìm attributes (private fields)
        attr_pattern = r'private\s+(?:static\s+)?(?:final\s+)?(\w+(?:<.*?>)?)\s+(\w+)\s*[;=]'
        attributes = re.findall(attr_pattern, content)
        class_info.attributes = [f"{attr_type} {attr_name}" for attr_type, attr_name in attributes]
        
        # Tìm methods (public methods)
        method_pattern = r'public\s+(?:static\s+)?(?:<.*?>\s+)?(\w+(?:<.*?>)?)\s+(\w+)\s*\([^)]*\)'
        methods = re.findall(method_pattern, content)
        class_info.methods = [f"{ret_type} {method_name}()" for ret_type, method_name in methods]
        
        # Phân tích responsibilities từ comments
        comment_pattern = r'/\*\*.*?\*/'
        comments = re.findall(comment_pattern, content, re.DOTALL)
        for comment in comments:
            if 'Trách nhiệm' in comment or 'Responsibility' in comment:
                lines = comment.split('\n')
                for line in lines:
                    if 'Trách nhiệm' in line or 'Responsibility' in line:
                        desc = re.sub(r'[*/]', '', line).strip()
                        if desc:
                            class_info.responsibilities.append(desc)
        
        # Tìm collaborators từ imports và usage
        for imp in class_info.imports:
            parts = imp.split('.')
            if len(parts) > 0:
                last_part = parts[-1]
                if last_part.endswith('Service') or last_part.endswith('Model') or \
                   last_part in ['DatabaseConnector', 'Helper', 'AccessManager', 'UserRole']:
                    class_info.collaborators.add(last_part)
        
        return class_info
    
    def scan_project(self):
        """Quét toàn bộ project để tìm các file Java"""
        java_files = list(self.base_path.rglob('*.java'))
        
        for java_file in java_files:
            # Bỏ qua test files
            if 'test' in str(java_file).lower():
                continue
            
            class_info = self.parse_file(java_file)
            if class_info:
                full_name = f"{class_info.package}.{class_info.name}"
                self.classes[full_name] = class_info
        
        # Phân tích thêm collaborators từ usage
        self._analyze_collaborators()
    
    def _analyze_collaborators(self):
        """Phân tích thêm collaborators từ việc sử dụng các class khác"""
        for class_info in self.classes.values():
            file_path = self.base_path / class_info.file_path
            if not file_path.exists():
                continue
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Tìm usage của các class khác
            for other_class in self.classes.values():
                if other_class.name != class_info.name:
                    # Tìm trong code
                    pattern = rf'\b{other_class.name}\b'
                    if re.search(pattern, content):
                        class_info.collaborators.add(other_class.name)

class CRCGenerator:
    """Generator để tạo thiết kế CRC"""
    
    def __init__(self, parser: JavaParser):
        self.parser = parser
    
    def generate_markdown(self, output_file: str = "CRC_Design_Auto.md"):
        """Tạo file markdown chứa thiết kế CRC"""
        output_path = self.parser.base_path / output_file
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# THIẾT KẾ CRC TỰ ĐỘNG\n")
            f.write("## Hệ thống Quản lý Chung cư BlueMoon\n")
            f.write("*(File này được tạo tự động bởi generate_crc.py)*\n\n")
            f.write("---\n\n")
            
            # Nhóm theo package
            packages = defaultdict(list)
            for class_info in self.parser.classes.values():
                packages[class_info.package].append(class_info)
            
            for package_name in sorted(packages.keys()):
                f.write(f"## PACKAGE: {package_name}\n\n")
                
                for class_info in sorted(packages[package_name], key=lambda x: x.name):
                    self._write_class_crc(f, class_info)
                    f.write("\n---\n\n")
            
            # Viết quan hệ giữa các lớp
            f.write("## QUAN HỆ GIỮA CÁC LỚP\n\n")
            self._write_relationships(f)
    
    def _write_class_crc(self, f, class_info: ClassInfo):
        """Viết thẻ CRC cho một lớp"""
        f.write(f"### Class: {class_info.name}\n")
        f.write(f"**Loại:** {class_info.type}\n")
        f.write(f"**Package:** {class_info.package}\n")
        f.write(f"**File:** {class_info.file_path}\n\n")
        
        # Responsibilities
        f.write("**Trách nhiệm (Responsibilities):**\n")
        if class_info.responsibilities:
            for resp in class_info.responsibilities:
                f.write(f"- {resp}\n")
        else:
            # Tự động tạo từ type và name
            if class_info.type == 'interface':
                f.write(f"- Định nghĩa contract cho các nghiệp vụ liên quan đến {class_info.name}\n")
            elif 'Service' in class_info.name:
                f.write(f"- Triển khai logic nghiệp vụ quản lý {class_info.name.replace('Service', '')}\n")
            elif 'Model' in class_info.name or class_info.name in ['HoGiaDinh', 'NhanKhau', 'PhieuThu']:
                f.write(f"- Lưu trữ thông tin về {class_info.name}\n")
            else:
                f.write(f"- Quản lý {class_info.name}\n")
        f.write("\n")
        
        # Collaborators
        f.write("**Cộng tác (Collaborators):**\n")
        if class_info.collaborators:
            for collab in sorted(class_info.collaborators):
                f.write(f"- {collab}\n")
        else:
            f.write("- (Chưa xác định)\n")
        f.write("\n")
        
        # Attributes
        f.write("**Thuộc tính (Attributes):**\n")
        if class_info.attributes:
            for attr in class_info.attributes[:10]:  # Giới hạn 10 thuộc tính đầu
                f.write(f"- {attr}\n")
            if len(class_info.attributes) > 10:
                f.write(f"- ... và {len(class_info.attributes) - 10} thuộc tính khác\n")
        else:
            f.write("- (Không có thuộc tính private)\n")
        f.write("\n")
        
        # Methods
        if class_info.methods:
            f.write("**Methods (một số):**\n")
            for method in class_info.methods[:10]:  # Giới hạn 10 methods đầu
                f.write(f"- {method}\n")
            if len(class_info.methods) > 10:
                f.write(f"- ... và {len(class_info.methods) - 10} methods khác\n")
            f.write("\n")
    
    def _write_relationships(self, f):
        """Viết quan hệ giữa các lớp"""
        # Implementation relationships
        implementations = []
        for class_info in self.parser.classes.values():
            if class_info.type == 'interface':
                # Tìm implementation
                impl_name = class_info.name.replace('Service', 'ServiceImpl')
                for other in self.parser.classes.values():
                    if other.name == impl_name:
                        implementations.append((class_info.name, other.name))
        
        if implementations:
            f.write("### Quan hệ Implementation:\n")
            for interface, impl in implementations:
                f.write(f"- `{impl}` implements `{interface}`\n")
            f.write("\n")
        
        # Dependency relationships
        f.write("### Quan hệ Dependency:\n")
        for class_info in self.parser.classes.values():
            if class_info.collaborators:
                for collab in sorted(class_info.collaborators):
                    f.write(f"- `{class_info.name}` → `{collab}`\n")
        f.write("\n")
    
    def generate_html_table(self, output_file: str = "CRC_Table.html"):
        """Tạo bảng CRC dạng HTML"""
        output_path = self.parser.base_path / output_file
        
        html = """<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CRC Cards - BlueMoon</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 20px;
            background-color: #f5f5f5;
        }
        .crc-container {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(400px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }
        .crc-card {
            background: white;
            border: 2px solid #333;
            border-radius: 8px;
            padding: 15px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .crc-card h3 {
            margin-top: 0;
            color: #0066cc;
            border-bottom: 2px solid #0066cc;
            padding-bottom: 5px;
        }
        .crc-section {
            margin: 10px 0;
        }
        .crc-section h4 {
            color: #666;
            font-size: 14px;
            margin-bottom: 5px;
        }
        .crc-section ul {
            margin: 5px 0;
            padding-left: 20px;
        }
        .crc-section li {
            margin: 3px 0;
            font-size: 12px;
        }
        .package-header {
            background: #0066cc;
            color: white;
            padding: 10px;
            margin: 20px 0 10px 0;
            border-radius: 5px;
        }
    </style>
</head>
<body>
    <h1>Thiết kế CRC - Hệ thống Quản lý Chung cư BlueMoon</h1>
    <p><em>File được tạo tự động bởi generate_crc.py</em></p>
"""
        
        # Nhóm theo package
        packages = defaultdict(list)
        for class_info in self.parser.classes.values():
            packages[class_info.package].append(class_info)
        
        for package_name in sorted(packages.keys()):
            html += f'    <div class="package-header"><h2>Package: {package_name}</h2></div>\n'
            html += '    <div class="crc-container">\n'
            
            for class_info in sorted(packages[package_name], key=lambda x: x.name):
                html += self._generate_card_html(class_info)
            
            html += '    </div>\n'
        
        html += """</body>
</html>"""
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)
    
    def generate_word_document(self, output_file: str = "CRC_Design.docx"):
        """Tạo file Word chứa các bảng CRC"""
        if not DOCX_AVAILABLE:
            print("Không thể tạo file Word: python-docx chưa được cài đặt")
            print("Chạy: pip install python-docx")
            return
        
        output_path = self.parser.base_path / output_file
        
        # Tạo document mới
        doc = Document()
        
        # Thiết lập font mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Tiêu đề chính
        title = doc.add_heading('THIẾT KẾ CRC', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Hệ thống Quản lý Chung cư BlueMoon')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.bold = True
        
        note = doc.add_paragraph('(File này được tạo tự động bởi generate_crc.py)')
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_format = note.runs[0].font
        note_format.italic = True
        note_format.size = Pt(10)
        
        doc.add_paragraph()  # Dòng trống
        
        # Nhóm theo package
        packages = defaultdict(list)
        for class_info in self.parser.classes.values():
            packages[class_info.package].append(class_info)
        
        for package_name in sorted(packages.keys()):
            # Tiêu đề package
            package_heading = doc.add_heading(f'Package: {package_name}', 1)
            package_heading_format = package_heading.runs[0].font
            package_heading_format.color.rgb = RGBColor(0, 102, 204)
            
            # Tạo bảng CRC cho các lớp trong package
            for class_info in sorted(packages[package_name], key=lambda x: x.name):
                self._create_crc_table(doc, class_info)
                # Thêm bảng vàng UML diagram sau mỗi bảng CRC
                self._create_uml_diagram_table(doc, class_info)
                doc.add_paragraph()  # Dòng trống giữa các bảng
        
        # Lưu file
        doc.save(output_path)
    
    def _create_crc_table(self, doc, class_info: ClassInfo):
        """Tạo bảng CRC cho một lớp trong Word document với header màu vàng"""
        # Tạo bảng với 2 cột: Class name và Responsibilities/Collaborators
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row - màu vàng/cam
        header_cell = table.rows[0].cells[0]
        header_cell.merge(table.rows[0].cells[1])  # Merge 2 cột cho header
        header_cell.width = Inches(8)
        
        # Thiết lập màu nền vàng/cam cho header
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(r'<w:shd {} w:fill="FFD700"/>'.format(nsdecls('w')))
        header_cell._element.get_or_add_tcPr().append(shading)
        
        # Tên lớp trong header
        header_para = header_cell.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(f'Class {class_info.name}')
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen
        
        # Row 2 - Cột 1: Mô tả và thông tin
        cell1 = table.rows[1].cells[0]
        cell1.width = Inches(2.5)
        
        # Mô tả trách nhiệm ngắn gọn
        desc_para = cell1.paragraphs[0]
        auto_resp = self._generate_responsibility(class_info)
        desc_run = desc_para.add_run(auto_resp)
        desc_run.font.name = 'Times New Roman'
        desc_run.font.size = Pt(11)
        
        # Thông tin loại và package
        info_para = cell1.add_paragraph()
        info_run = info_para.add_run(f'({class_info.type})')
        info_run.font.size = Pt(9)
        info_run.font.italic = True
        
        # Row 2 - Cột 2: Thuộc tính đầy đủ
        cell2 = table.rows[1].cells[1]
        cell2.width = Inches(5.5)
        
        # Thuộc tính - liệt kê đầy đủ
        if class_info.attributes:
            attr_para = cell2.paragraphs[0]
            # Liệt kê TẤT CẢ thuộc tính, không giới hạn
            for attr in class_info.attributes:
                # Format: private Type name;
                attr_text = attr_para.add_run(f'{attr};\n')
                attr_text.font.name = 'Courier New'  # Font monospace cho code
                attr_text.font.size = Pt(10)
        else:
            attr_para = cell2.paragraphs[0]
            no_attr = attr_para.add_run('(Không có thuộc tính private)')
            no_attr.font.size = Pt(10)
            no_attr.font.italic = True
    
    def _create_uml_diagram_table(self, doc, class_info: ClassInfo):
        """Tạo bảng vàng UML diagram cho một lớp (giống như trong hình)"""
        # Tạo bảng 1 cột với nền vàng
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Light Grid Accent 1'
        
        # Cell duy nhất với nền vàng
        cell = table.rows[0].cells[0]
        cell.width = Inches(8)
        
        # Thiết lập màu nền vàng cho toàn bộ bảng
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(r'<w:shd {} w:fill="FFFF99"/>'.format(nsdecls('w')))  # Màu vàng nhạt
        cell._element.get_or_add_tcPr().append(shading)
        
        # Tên lớp (UML style)
        class_para = cell.paragraphs[0]
        class_run = class_para.add_run(f'{class_info.name}\n')
        class_run.font.name = 'Times New Roman'
        class_run.font.size = Pt(14)
        class_run.font.bold = True
        class_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thuộc tính dạng UML (- attribute : Type)
        if class_info.attributes:
            for attr in class_info.attributes:
                # Parse attribute: "Type name" -> "- name : Type"
                attr_parts = attr.strip().split()
                if len(attr_parts) >= 2:
                    attr_type = attr_parts[0]
                    attr_name = attr_parts[1]
                    # Format UML: - name : Type
                    uml_attr = cell.add_paragraph()
                    uml_run = uml_attr.add_run(f'- {attr_name} : {attr_type}')
                    uml_run.font.name = 'Courier New'
                    uml_run.font.size = Pt(11)
                else:
                    # Fallback nếu không parse được
                    uml_attr = cell.add_paragraph()
                    uml_run = uml_attr.add_run(f'- {attr}')
                    uml_run.font.name = 'Courier New'
                    uml_run.font.size = Pt(11)
        else:
            no_attr_para = cell.add_paragraph()
            no_attr_run = no_attr_para.add_run('(Không có thuộc tính)')
            no_attr_run.font.size = Pt(10)
            no_attr_run.font.italic = True
    
    def _generate_responsibility(self, class_info: ClassInfo) -> str:
        """Tự động tạo responsibility cho lớp"""
        if class_info.type == 'interface':
            return f'Định nghĩa contract cho các nghiệp vụ liên quan đến {class_info.name}'
        elif 'Service' in class_info.name:
            domain = class_info.name.replace('Service', '').replace('Impl', '')
            return f'Triển khai logic nghiệp vụ quản lý {domain}'
        elif class_info.name in ['HoGiaDinh', 'NhanKhau', 'PhieuThu', 'ChiTietThu', 'KhoanThu', 
                                  'DotThu', 'TaiKhoan', 'Phong', 'PhuongTien', 'LichSuNhanKhau']:
            return f'Lưu trữ thông tin về {class_info.name}'
        elif class_info.name == 'WebServer':
            return 'Xử lý HTTP requests và routing, chuyển đổi JSON'
        elif class_info.name == 'DatabaseConnector':
            return 'Quản lý kết nối đến PostgreSQL database'
        elif class_info.name == 'Helper':
            return 'Cung cấp các utility methods hỗ trợ (hash password, verify password)'
        elif class_info.name == 'UserRole':
            return 'Định nghĩa các vai trò người dùng trong hệ thống'
        elif class_info.name == 'AccessManager':
            return 'Quản lý quyền truy cập dựa trên vai trò người dùng'
        else:
            return f'Quản lý {class_info.name}'
    
    def _generate_card_html(self, class_info: ClassInfo) -> str:
        """Tạo HTML cho một thẻ CRC"""
        html = f"""        <div class="crc-card">
            <h3>{class_info.name} ({class_info.type})</h3>
            
            <div class="crc-section">
                <h4>Trách nhiệm (Responsibilities):</h4>
                <ul>
"""
        
        if class_info.responsibilities:
            for resp in class_info.responsibilities[:5]:
                html += f"                    <li>{resp}</li>\n"
        else:
            html += f"                    <li>Quản lý {class_info.name}</li>\n"
        
        html += """                </ul>
            </div>
            
            <div class="crc-section">
                <h4>Cộng tác (Collaborators):</h4>
                <ul>
"""
        
        if class_info.collaborators:
            for collab in sorted(list(class_info.collaborators))[:8]:
                html += f"                    <li>{collab}</li>\n"
        else:
            html += "                    <li>(Chưa xác định)</li>\n"
        
        html += """                </ul>
            </div>
            
            <div class="crc-section">
                <h4>Thuộc tính (Attributes):</h4>
                <ul>
"""
        
        if class_info.attributes:
            for attr in class_info.attributes[:5]:
                html += f"                    <li>{attr}</li>\n"
            if len(class_info.attributes) > 5:
                html += f"                    <li><em>... và {len(class_info.attributes) - 5} thuộc tính khác</em></li>\n"
        else:
            html += "                    <li>(Không có)</li>\n"
        
        html += """                </ul>
            </div>
        </div>
"""
        return html

def main():
    """Hàm main"""
    # Đường dẫn đến thư mục src/main/java/com/bluemoon
    base_path = Path(__file__).parent / "src" / "main" / "java" / "com" / "bluemoon"
    
    if not base_path.exists():
        print(f"Không tìm thấy thư mục: {base_path}")
        print("Vui lòng chạy script từ thư mục gốc của project")
        return
    
    print("Đang quét các file Java...")
    parser = JavaParser(base_path)
    parser.scan_project()
    
    print(f"Đã tìm thấy {len(parser.classes)} lớp")
    
    print("Đang tạo thiết kế CRC...")
    generator = CRCGenerator(parser)
    
    # Tạo file markdown
    generator.generate_markdown("CRC_Design_Auto.md")
    print("✓ Đã tạo file: CRC_Design_Auto.md")
    
    # Tạo file HTML
    generator.generate_html_table("CRC_Table.html")
    print("✓ Đã tạo file: CRC_Table.html")
    
    # Tạo file Word
    generator.generate_word_document("CRC_Design.docx")
    if DOCX_AVAILABLE:
        print("✓ Đã tạo file: CRC_Design.docx")
    else:
        print("⚠ Không thể tạo file Word (cần cài: pip install python-docx)")
    
    print("\nHoàn thành!")

if __name__ == "__main__":
    main()

